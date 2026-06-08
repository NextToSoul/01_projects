import docx
path = r"E:/Codex Workspace/01_Projects/项目文档/通讯协议/立即遥控指令与遥测应答.docx"
doc = docx.Document(path)
print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"[{i}]",t)
print()
print("=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print()
    print("--- Table",ti,"(",len(table.rows),"rows x",len(table.columns),"cols) ---")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip().replace(chr(10)," | ") for c in row.cells]
        print("  Row",ri,":"," | ".join(cells))
